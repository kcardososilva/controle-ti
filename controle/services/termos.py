# -*- coding: utf-8 -*-
"""
Geração dos Termos de Responsabilidade (Entrega e Devolução).

O documento é SEMPRE gerado a partir dos modelos base que ficam em
``docs_templates/termos/`` — assim, qualquer atualização feita no documento base
(texto, formatação, novas linhas) passa a valer automaticamente no termo gerado.

O preenchimento localiza cada linha da tabela pelo RÓTULO da 1ª célula (não por
índice fixo) e é robusto contra:
  - linhas duplicadas (cabeçalho + digitação) — todas as ocorrências do mesmo
    campo são tratadas, evitando que dados de exemplo do modelo permaneçam;
  - células mescladas (ex.: bloco de "Observações:") — o rótulo é preservado e o
    conteúdo residual é limpo.
"""

import re
import unicodedata
from io import BytesIO
from pathlib import Path

from django.conf import settings
from django.utils import timezone
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt
from ProjetoEstoque.models import MovimentacaoItem


TERMOS_DIR = Path(settings.BASE_DIR) / "docs_templates" / "termos"
TEMPLATE_ENTREGA = TERMOS_DIR / "TEMPLATE BRANCO - Termo_de_Entrega.docx"
TEMPLATE_DEVOLUCAO = TERMOS_DIR / "TEMPLATE BRANCO - Termo_de_Devolução.docx"


def _safe(value, default="-"):
    if value is None:
        return default
    value = str(value).strip()
    return value if value else default


def _normalizar(texto):
    """Minúsculas sem acentos — para comparar rótulos do template de forma robusta."""
    texto = (texto or "").strip().lower()
    return "".join(
        c for c in unicodedata.normalize("NFD", texto)
        if unicodedata.category(c) != "Mn"
    )


def _numero_termo_auto(item, colaborador, nome_colaborador):
    """
    Numeração automática do termo no formato pedido pelo TI:
        {Nº DE SÉRIE} - {NOME DO SOLICITANTE} - {CENTRO DE CUSTO}

    O nome do centro de custo vem por último. Usa o centro de custo do
    colaborador e, na ausência dele, o do próprio equipamento.
    """
    partes = []

    serie = _safe(getattr(item, "numero_serie", None), "")
    if serie and serie != "-":
        partes.append(serie)

    if nome_colaborador and nome_colaborador != "-":
        partes.append(nome_colaborador)

    cc_obj = getattr(colaborador, "centro_custo", None) or getattr(item, "centro_custo", None)
    cc_nome = _safe(getattr(cc_obj, "departamento", None), "")
    if not cc_nome or cc_nome == "-":
        cc_nome = _safe(getattr(cc_obj, "numero", None), "")
    if cc_nome and cc_nome != "-":
        partes.append(cc_nome)

    return " - ".join(partes)


def _clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def _replace_paragraph_text(paragraph, text):
    style = paragraph.style
    _clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    paragraph.style = style

    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(4)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.15

    return run


def _find_paragraph_contains(doc, needle):
    for p in doc.paragraphs:
        if needle in (p.text or ""):
            return p
    return None


def _set_cell(cell, text):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(text)


def _usuario_display(usuario):
    if not usuario:
        return "-"

    return _safe(getattr(usuario, "nome", None), "-")


def get_usuario_atual_item(item):
    """
    Retorna o usuário atualmente com o item: última entrega (transferencia + tipo_transferencia=entrega)
    sem uma devolução posterior. Garante que devolução não seja confundida com posse atual.
    """
    from ProjetoEstoque.models import TipoMovimentacaoChoices, TipoTransferenciaChoices

    ultima_entrega = (
        MovimentacaoItem.objects
        .filter(
            item=item,
            tipo_movimentacao=TipoMovimentacaoChoices.TRANSFERENCIA,
            tipo_transferencia=TipoTransferenciaChoices.ENTREGA,
            usuario__isnull=False,
        )
        .select_related("usuario")
        .order_by("-created_at")
        .first()
    )

    if not ultima_entrega:
        return None

    # Verifica se houve devolução posterior
    devolucao_posterior = (
        MovimentacaoItem.objects
        .filter(
            item=item,
            tipo_movimentacao=TipoMovimentacaoChoices.TRANSFERENCIA,
            tipo_transferencia=TipoTransferenciaChoices.DEVOLUCAO,
            created_at__gt=ultima_entrega.created_at,
        )
        .exists()
    )

    if devolucao_posterior:
        return None

    return ultima_entrega.usuario


def _build_estabelecimento_line(estabelecimento):
    checks = {
        "rio_do_meio": " ",
        "karitel": " ",
        "sao_paulo": " ",
        "sta_edwiges": " ",
    }
    if estabelecimento in checks:
        checks[estabelecimento] = "X"

    return (
        f"({checks['rio_do_meio']}) Rio do Meio   "
        f"({checks['karitel']}) Karitel   "
        f"({checks['sao_paulo']}) São Paulo   "
        f"({checks['sta_edwiges']}) Sta.Edwiges"
    )


def _build_dados(item, tipo, form_data):
    hoje = timezone.localdate()
    colaborador = form_data.get("colaborador")
    nome_colaborador = _usuario_display(colaborador)
    email_colaborador = _safe(getattr(colaborador, "email", None), "")
    centro_custo_colaborador = _safe(getattr(getattr(colaborador, "centro_custo", None), "departamento", None), "")
    funcao_colaborador = _safe(getattr(getattr(colaborador, "funcao", None), "nome", None), "")
    localidade_colaborador = _safe(getattr(getattr(colaborador, "localidade", None), "local", None), "")

    descricao = (
        f"{_safe(item.nome)} | "
        f"Modelo: {_safe(item.modelo)} | "
        f"Marca: {_safe(item.marca)} | "
        f"Centro de Custo: {_safe(getattr(item.centro_custo, 'departamento', None))}"
    )

    # Numeração: {Nº de série} - {nome do solicitante} - {centro de custo}.
    # O campo do formulário continua válido como sobrescrita manual.
    numero_termo = _safe(form_data.get("numero_termo"), "")
    if not numero_termo:
        numero_termo = _numero_termo_auto(item, colaborador, nome_colaborador)

    return {
        "tipo": tipo,
        "data_hoje": hoje.strftime("%d/%m/%Y"),
        "numero_termo": numero_termo,
        "numero_chamado": _safe(form_data.get("numero_chamado"), ""),
        "nome_colaborador": nome_colaborador,
        "email_colaborador": email_colaborador,
        "centro_custo_colaborador": centro_custo_colaborador,
        "funcao_colaborador": funcao_colaborador,
        "localidade_colaborador": localidade_colaborador,
        "descricao_equipamento": descricao,
        "serie": _safe(item.numero_serie),
        "acessorios": _safe(form_data.get("acessorios"), ""),
        "plaqueta": _safe(getattr(item, "patrimonio", None), ""),
        "estabelecimento": _build_estabelecimento_line(form_data.get("estabelecimento")),
        "observacoes": _safe(form_data.get("observacoes"), ""),
        "responsavel_ti_nome": _safe(form_data.get("responsavel_ti_nome"), ""),
    }


def _fill_intro(doc, dados):
    p_intro = _find_paragraph_contains(doc, "Eu, __")
    if not p_intro:
        return

    if dados["tipo"] == "entrega":
        texto = (
            f"Eu, {dados['nome_colaborador']}, portador(a) de cédula de identidade RG nº "
            f"_______________________________________________ e inscrito no CPF/MF nº "
            f"_____________________________ de agora em diante chamado de USUÁRIO, "
            f"declaro que recebi, li e entendi as obrigações desse Termo de Responsabilidade "
            f"pela Guarda e Uso de Equipamentos Corporativos da SANTA COLOMBA (Termo) e me "
            f"comprometo, de forma irretratável, com as regras e obrigações abaixo:"
        )
    else:
        texto = (
            f"Eu, {dados['nome_colaborador']}, portador(a) de cédula de identidade RG nº "
            f"_______________________________________________ e inscrito no CPF/MF nº "
            f"_____________________________ declaro que devolvi o(s) equipamento(s) descrito(s) "
            f"nesse documento, pertencente(s) à SANTA COLOMBA, nas condições especificadas. "
            f"Declaro estar ciente de que, caso haja danos não reportados ou perda de algum item, "
            f"poderei ser responsabilizado conforme as normas internas da empresa."
        )

    _replace_paragraph_text(p_intro, texto)


def _fill_signatures(doc, dados):
    nomes_encontrados = 0
    datas_encontradas = 0

    for p in doc.paragraphs:
        txt = (p.text or "").strip()

        if txt == "Colaborador:":
            _replace_paragraph_text(p, f"Colaborador: {dados['nome_colaborador']}")

        elif txt == "Nome:":
            nomes_encontrados += 1
            if nomes_encontrados == 1:
                _replace_paragraph_text(p, f"Nome: {dados['nome_colaborador']}")
            else:
                _replace_paragraph_text(p, f"Nome: {dados['responsavel_ti_nome']}")

        elif txt == "Data:":
            datas_encontradas += 1
            _replace_paragraph_text(p, f"Data: {dados['data_hoje']}")


def _fill_main_table(doc, dados):
    """
    Preenche a tabela do termo localizando cada linha pelo RÓTULO da 1ª célula.

    Robustez contra dados de exemplo do modelo:
      - cada CAMPO é preenchido apenas na 1ª linha correspondente; as demais
        linhas com o mesmo rótulo têm o valor LIMPO (remove exemplos residuais);
      - células mescladas (rótulo + valor juntos, como o bloco "Observações:")
        são reescritas preservando o rótulo;
      - células físicas mescladas são tratadas uma única vez.
    """
    if not doc.tables:
        return

    table = doc.tables[0]

    # ordem importa: o predicado de "termo" exclui linhas de "chamado".
    campos = [
        (lambda t: "termo" in t and "chamado" not in t, dados["numero_termo"]),
        (lambda t: "chamado" in t, dados["numero_chamado"]),
        (lambda t: "descric" in t, dados["descricao_equipamento"]),
        (lambda t: t.startswith("serie"), dados["serie"]),
        (lambda t: "acessorio" in t, dados["acessorios"]),
        (lambda t: "plaqueta" in t, dados["plaqueta"]),
        (lambda t: "estabelecimento" in t, dados["estabelecimento"]),
        (lambda t: "observac" in t, dados["observacoes"]),
    ]

    campo_preenchido = set()     # índices de campos cuja 1ª linha já recebeu valor
    celulas_tratadas = []        # refs de <w:tc> já processadas (dedup de mesclagem)

    def _ja_tratada(tc):
        return any(tc is x for x in celulas_tratadas)

    def _vazio(v):
        return (not v) or str(v).strip() in ("", "-")

    for row in table.rows:
        cells = row.cells
        if not cells:
            continue

        label_raw = cells[0].text or ""
        label = _normalizar(label_raw)
        if not label:
            continue

        idx = next((i for i, (pred, _) in enumerate(campos) if pred(label)), None)
        if idx is None:
            continue

        valor = campos[idx][1]
        # rótulo e valor na mesma célula física? (bloco mesclado de observações)
        mesclada = len(cells) > 1 and cells[1]._tc is cells[0]._tc
        c_val = cells[0] if mesclada else (cells[1] if len(cells) > 1 else cells[0])

        if _ja_tratada(c_val._tc):
            continue

        primeira = idx not in campo_preenchido

        if mesclada:
            base = label_raw.split(":")[0].strip()
            if primeira and not _vazio(valor):
                _set_cell(c_val, f"{base}: {valor}")
            elif primeira:
                _set_cell(c_val, f"{base}:")
            else:
                _set_cell(c_val, "")
        else:
            _set_cell(c_val, "" if (_vazio(valor) or not primeira) else str(valor))

        campo_preenchido.add(idx)
        celulas_tratadas.append(c_val._tc)


def _ajustar_layout_documento(doc):
    """
    Ajuste de layout NÃO destrutivo: preserva a formatação do modelo base e só
    aplica margens A4, centraliza o título e garante uma fonte mínima onde o run
    não define nenhuma. Não sobrescreve a formatação do template.
    """
    for section in doc.sections:
        section.top_margin = 720000      # ~2 cm
        section.bottom_margin = 720000   # ~2 cm
        section.left_margin = 720000     # ~2 cm
        section.right_margin = 720000    # ~2 cm

    for p in doc.paragraphs:
        texto = (p.text or "").strip()

        if texto.startswith("TERMO DE RESPONSABILIDADE") or texto.startswith("POLÍTICA"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in p.runs:
            if run.font.size is None:
                run.font.size = Pt(10.5)


def gerar_termo_docx(item, tipo, form_data):
    if tipo not in ("entrega", "devolucao"):
        raise ValueError("Tipo inválido para termo.")

    template_path = TEMPLATE_ENTREGA if tipo == "entrega" else TEMPLATE_DEVOLUCAO
    if not template_path.exists():
        raise FileNotFoundError(f"Template não encontrado: {template_path}")

    doc = Document(str(template_path))
    dados = _build_dados(item, tipo, form_data)

    _fill_intro(doc, dados)
    _fill_main_table(doc, dados)
    _fill_signatures(doc, dados)
    _ajustar_layout_documento(doc)

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    nome_item = _safe(item.nome, "equipamento").replace(" ", "_")
    nome_arquivo = f"termo_{tipo}_{item.pk}_{nome_item}.docx"

    return output, nome_arquivo
